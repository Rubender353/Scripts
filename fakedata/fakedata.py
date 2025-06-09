import os
import random
import string
from faker import Faker
from openpyxl import Workbook
from docx import Document
from reportlab.pdfgen import canvas

fake = Faker()

# Directory to store files
output_dir = "generated_files"
os.makedirs(output_dir, exist_ok=True)

# Sample filenames and departments
departments = {
    "Finance": ["Quarterly_Report", "Budget_Projection", "Expense_Overview"],
    "Payroll": ["Salary_Distribution", "Tax_Document", "Payroll_Summary"],
    "Administration": ["Meeting_Minutes", "Company_Policy", "Staff_Directory"],
    "IT": ["System_Log", "Asset_Inventory", "Software_Licenses"]
}

# Helper function to generate a random file size up to 20MB
def get_target_size_bytes():
    return random.randint(1_000_000, 20_000_000)  # 100 KB to 20MB

def generate_excel(file_path, target_size):
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"

    # Add initial data to create the file
    for _ in range(100):
        data = [fake.name(), fake.company(), fake.email(), fake.date(), fake.random_int()]
        ws.append(data)

    wb.save(file_path)

def generate_word(file_path, target_size):
    doc = Document()
    while True:
        doc.add_heading(fake.bs().title(), level=2)
        for _ in range(10):
            doc.add_paragraph(fake.paragraph(nb_sentences=5))
        doc.save(file_path)
        if os.path.getsize(file_path) >= target_size:
            break

def generate_pdf(file_path, target_size):
    c = canvas.Canvas(file_path)
    text = c.beginText(40, 800)
    while True:
        for _ in range(50):
            text.textLine(fake.sentence())
        c.drawText(text)
        c.showPage()
        c.save()
        if os.path.getsize(file_path) >= target_size:
            break
        text = c.beginText(40, 800)

# Generate files
file_count = 5  # You can increase this
for _ in range(file_count):
    dept = random.choice(list(departments.keys()))
    title = random.choice(departments[dept])
    ext = random.choice([".xlsx", ".docx", ".pdf"])
    filename = f"{dept}_{title}_{fake.random_int(100, 999)}{ext}"
    file_path = os.path.join(output_dir, filename)
    target_size = get_target_size_bytes()

    print(f"Generating {filename} with target size {target_size / 1_000_000:.2f} MB...")

    if ext == ".xlsx":
        generate_excel(file_path, target_size)
    elif ext == ".docx":
        generate_word(file_path, target_size)
    elif ext == ".pdf":
        generate_pdf(file_path, target_size)

print("✅ File generation complete.")
